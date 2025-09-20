"""
FastAPI application for the AI Research Agent
"""
from fastapi import FastAPI, HTTPException, BackgroundTasks
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.exceptions import RequestValidationError
from pydantic import BaseModel, Field
from typing import List, Optional, Dict, Any
import asyncio
import logging
from datetime import datetime, timezone
from starlette.requests import Request
from io import BytesIO

from agent import AIResearchAgent
from models import ResearchRequest, ResearchStatus
from database import db_manager

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
# Celery imports removed - using synchronous processing for now

app = FastAPI(
    title="AI Research Agent",
    description="An AI-powered research agent that accepts topics and returns structured research results",
    version="1.0.0"
)

# Add CORS middleware
import os
ALLOWED_ORIGINS = os.getenv("ALLOWED_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000").split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["*"],
)

# Initialize the agent
research_agent = AIResearchAgent()

# Job status tracking and task handles
job_status = {}
_tasks: dict[str, asyncio.Task] = {}

async def _run_research_async(research_id: str, topic: str):
    try:
        if research_id in job_status:
            job_status[research_id]["status"] = "IN_PROGRESS"
        # run workflow with fixed id so DB record matches polled id
        await research_agent.research_topic(topic, research_id=research_id)
        # give DB a brief moment to flush visibility
        await asyncio.sleep(0.1)
        if research_id in job_status and job_status[research_id].get("status") != "CANCELLED":
            job_status[research_id]["status"] = "COMPLETED"
            job_status[research_id]["completed_at"] = datetime.now(timezone.utc).isoformat()
    except asyncio.CancelledError:
        if research_id in job_status:
            job_status[research_id]["status"] = "CANCELLED"
            job_status[research_id]["completed_at"] = datetime.now(timezone.utc).isoformat()
        raise
    except Exception as e:
        if research_id in job_status:
            job_status[research_id]["status"] = "FAILED"
            job_status[research_id]["error"] = str(e)

# Exception handlers for clearer 4xx responses
@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    return JSONResponse(
        status_code=422,
        content={
            "error": "Invalid request payload",
            "path": request.url.path,
            "details": exc.errors(),
        },
    )

@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail if isinstance(exc.detail, str) else "HTTP error",
            "path": request.url.path,
        },
    )

# Request/Response models
class ResearchTopicRequest(BaseModel):
    topic: str = Field(..., min_length=3, description="Research topic (min 3 chars)")

class ResearchResponse(BaseModel):
    research_id: str
    status: str
    message: str
    estimated_completion_time: Optional[str] = None

class ResearchResultResponse(BaseModel):
    research_id: str
    topic: str
    status: str
    created_at: datetime
    completed_at: Optional[datetime]
    summary: Optional[str] = None
    key_findings: Optional[List[str]] = None
    sources: Optional[List[dict]] = None
    confidence_score: Optional[float] = None
    trace_log: List[str] = []
    steps: List[dict] = []
    results: Optional[Dict[str, Any]] = None  # Add results field

@app.get("/", response_class=HTMLResponse)
async def root():
    """
    Serve the main web interface
    """
    return """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>AI Research Agent</title>
        <style>
            body {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                max-width: 1200px;
                margin: 0 auto;
                padding: 20px;
                background-color: #f5f5f5;
            }
            .container {
                background: white;
                border-radius: 12px;
                padding: 30px;
                box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
            }
            h1 {
                color: #2c3e50;
                text-align: center;
                margin-bottom: 30px;
            }
            .form-group {
                margin-bottom: 20px;
            }
            label {
                display: block;
                margin-bottom: 8px;
                font-weight: 600;
                color: #34495e;
            }
            input[type="text"] {
                width: 100%;
                padding: 12px;
                border: 2px solid #e1e8ed;
                border-radius: 8px;
                font-size: 16px;
                transition: border-color 0.3s;
            }
            input[type="text"]:focus {
                outline: none;
                border-color: #3498db;
            }
            button {
                background: #3498db;
                color: white;
                border: none;
                padding: 12px 24px;
                border-radius: 8px;
                font-size: 16px;
                cursor: pointer;
                transition: background-color 0.3s;
            }
            button:hover {
                background: #2980b9;
            }
            button:disabled {
                background: #bdc3c7;
                cursor: not-allowed;
            }
            .results {
                margin-top: 30px;
                padding: 20px;
                background: #f8f9fa;
                border-radius: 8px;
                display: none;
            }
            .loading {
                text-align: center;
                color: #7f8c8d;
                font-style: italic;
            }
            .error {
                color: #e74c3c;
                background: #fdf2f2;
                padding: 10px;
                border-radius: 4px;
                margin: 10px 0;
            }
            .success {
                color: #27ae60;
                background: #f0f9f0;
                padding: 10px;
                border-radius: 4px;
                margin: 10px 0;
            }
            .research-item {
                background: white;
                border: 1px solid #e1e8ed;
                border-radius: 8px;
                padding: 15px;
                margin: 10px 0;
                cursor: pointer;
                transition: box-shadow 0.3s;
            }
            .research-item:hover {
                box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
            }
            .research-item h3 {
                margin: 0 0 10px 0;
                color: #2c3e50;
            }
            .research-item .status {
                display: inline-block;
                padding: 4px 8px;
                border-radius: 4px;
                font-size: 12px;
                font-weight: bold;
                text-transform: uppercase;
            }
            .status.completed {
                background: #d4edda;
                color: #155724;
            }
            .status.in-progress {
                background: #fff3cd;
                color: #856404;
            }
            .status.failed {
                background: #f8d7da;
                color: #721c24;
            }
            .status.pending {
                background: #e2e3e5;
                color: #383d41;
            }
            .trace-log {
                background: #f8f9fa;
                border: 1px solid #e9ecef;
                border-radius: 4px;
                padding: 10px;
                margin: 10px 0;
                font-family: monospace;
                font-size: 12px;
                max-height: 200px;
                overflow-y: auto;
            }
            .key-findings {
                margin: 15px 0;
            }
            .key-findings ul {
                margin: 0;
                padding-left: 20px;
            }
            .key-findings li {
                margin: 5px 0;
            }
            .sources {
                margin: 15px 0;
            }
            .source-item {
                background: #f8f9fa;
                border-left: 3px solid #3498db;
                padding: 10px;
                margin: 5px 0;
            }
            .source-item a {
                color: #3498db;
                text-decoration: none;
            }
            .source-item a:hover {
                text-decoration: underline;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>🔍 AI Research Agent</h1>
            <p style="text-align: center; color: #7f8c8d; margin-bottom: 30px;">
                Enter a topic and let our AI agent research it for you with explainable traces
            </p>
            
            <form id="researchForm">
                <div class="form-group">
                    <label for="topic">Research Topic:</label>
                    <input type="text" id="topic" name="topic" placeholder="e.g., Artificial Intelligence in Healthcare" required>
                </div>
                <button type="submit" id="submitBtn">Start Research</button>
            </form>
            
            <div id="results" class="results">
                <h2>Research Results</h2>
                <div id="researchList"></div>
            </div>
        </div>

        <script>
            const form = document.getElementById('researchForm');
            const results = document.getElementById('results');
            const researchList = document.getElementById('researchList');
            const submitBtn = document.getElementById('submitBtn');

            form.addEventListener('submit', async (e) => {
                e.preventDefault();
                const topic = document.getElementById('topic').value;
                
                if (!topic.trim()) {
                    alert('Please enter a research topic');
                    return;
                }

                submitBtn.disabled = true;
                submitBtn.textContent = 'Researching...';
                results.style.display = 'block';
                researchList.innerHTML = '<div class="loading">Starting research...</div>';

                try {
                    const response = await fetch('/research', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({ topic: topic })
                    });

                    const data = await response.json();
                    
                    if (response.ok) {
                        researchList.innerHTML = '<div class="success">Research started! Checking status...</div>';
                        pollResearchStatus(data.research_id);
                    } else {
                        researchList.innerHTML = `<div class="error">Error: ${data.detail}</div>`;
                    }
                } catch (error) {
                    researchList.innerHTML = `<div class="error">Error: ${error.message}</div>`;
                } finally {
                    submitBtn.disabled = false;
                    submitBtn.textContent = 'Start Research';
                }
            });

            async function pollResearchStatus(researchId) {
                const maxAttempts = 30; // 30 seconds max
                let attempts = 0;

                const poll = async () => {
                    try {
                        const response = await fetch(`/research/${researchId}`);
                        const data = await response.json();
                        
                        if (data.status === 'completed') {
                            displayResearchResult(data);
                        } else if (data.status === 'failed') {
                            researchList.innerHTML = `<div class="error">Research failed: ${data.trace_log[data.trace_log.length - 1] || 'Unknown error'}</div>`;
                        } else {
                            attempts++;
                            if (attempts < maxAttempts) {
                                researchList.innerHTML = `<div class="loading">Research in progress... (${attempts}/${maxAttempts})</div>`;
                                setTimeout(poll, 1000);
                            } else {
                                researchList.innerHTML = '<div class="error">Research timed out. Please try again.</div>';
                            }
                        }
                    } catch (error) {
                        researchList.innerHTML = `<div class="error">Error checking status: ${error.message}</div>`;
                    }
                };

                poll();
            }

            function displayResearchResult(data) {
                let html = `
                    <div class="research-item">
                        <h3>${data.topic}</h3>
                        <span class="status ${data.status}">${data.status}</span>
                        <p><strong>Completed:</strong> ${new Date(data.completed_at).toLocaleString()}</p>
                `;

                if (data.summary) {
                    html += `<div class="summary"><h4>Summary</h4><p>${data.summary.replace(/\\n/g, '<br>')}</p></div>`;
                }

                if (data.key_findings && data.key_findings.length > 0) {
                    html += `<div class="key-findings"><h4>Key Findings</h4><ul>`;
                    data.key_findings.forEach(finding => {
                        html += `<li>${finding}</li>`;
                    });
                    html += `</ul></div>`;
                }

                if (data.sources && data.sources.length > 0) {
                    html += `<div class="sources"><h4>Sources</h4>`;
                    data.sources.forEach(source => {
                        html += `
                            <div class="source-item">
                                <strong><a href="${source.url}" target="_blank">${source.title}</a></strong><br>
                                <small>${source.source} - ${source.snippet}</small>
                            </div>
                        `;
                    });
                    html += `</div>`;
                }

                if (data.confidence_score) {
                    html += `<p><strong>Confidence Score:</strong> ${(data.confidence_score * 100).toFixed(1)}%</p>`;
                }

                if (data.trace_log && data.trace_log.length > 0) {
                    html += `<div class="trace-log"><h4>Research Trace</h4><pre>${data.trace_log.join('\\n')}</pre></div>`;
                }

                html += `</div>`;
                researchList.innerHTML = html;
            }

            // Load existing research on page load
            window.addEventListener('load', async () => {
                try {
                    const response = await fetch('/research');
                    const data = await response.json();
                    
                    if (data.length > 0) {
                        results.style.display = 'block';
                        researchList.innerHTML = '<h3>Previous Research</h3>';
                        data.forEach(item => {
                            const itemDiv = document.createElement('div');
                            itemDiv.className = 'research-item';
                            itemDiv.innerHTML = `
                                <h3>${item.topic}</h3>
                                <span class="status ${item.status}">${item.status}</span>
                                <p><strong>Created:</strong> ${new Date(item.created_at).toLocaleString()}</p>
                                <button onclick="viewResearch('${item.research_id}')">View Details</button>
                            `;
                            researchList.appendChild(itemDiv);
                        });
                    }
                } catch (error) {
                    console.error('Error loading previous research:', error);
                }
            });

            async function viewResearch(researchId) {
                try {
                    const response = await fetch(`/research/${researchId}`);
                    const data = await response.json();
                    displayResearchResult(data);
                } catch (error) {
                    researchList.innerHTML = `<div class="error">Error loading research: ${error.message}</div>`;
                }
            }
        </script>
    </body>
    </html>
    """

@app.post("/research", response_model=ResearchResponse, status_code=202)
async def start_research(request: ResearchTopicRequest):
    """
    Start a new research session for the given topic using FastAPI async task
    """
    try:
        import uuid
        research_id = str(uuid.uuid4())

        # initialize job record
        job_status[research_id] = {
            "task_id": research_id,
            "status": "PENDING",
            "created_at": datetime.now(timezone.utc).isoformat(),
            "topic": request.topic
        }

        # schedule background task on event loop
        task = asyncio.create_task(_run_research_async(research_id, request.topic))
        _tasks[research_id] = task

        return ResearchResponse(
            research_id=research_id,
            status="PENDING",
            message="Research scheduled",
            estimated_completion_time="~2-3 minutes"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to start research: {str(e)}")

@app.post("/research/{research_id}/cancel")
async def cancel_research(research_id: str):
    task = _tasks.get(research_id)
    if not task:
        # If already finished or unknown
        if research_id in job_status:
            return {"cancelled": False, "status": job_status[research_id].get("status", "UNKNOWN")}
        raise HTTPException(status_code=404, detail="Research not found")
    if task.done() or task.cancelled():
        return {"cancelled": False, "status": job_status.get(research_id, {}).get("status", "COMPLETED")}
    task.cancel()
    job_status[research_id] = {**job_status.get(research_id, {}), "status": "CANCELLED"}
    return {"cancelled": True, "research_id": research_id}

@app.get("/research/{research_id}/status")
async def get_research_status(research_id: str):
    if research_id not in job_status:
        raise HTTPException(status_code=404, detail="Research job not found")

    # Compute progress from DB (steps out of 5) and provide snapshot when ready
    req = db_manager.get_research_request(research_id)
    current = len(req.steps) if req else 0
    total = 5

    info = job_status[research_id]
    
    # Enhanced progress tracking with step details
    progress_details = []
    if req and req.steps:
        for i, step in enumerate(req.steps):
            progress_details.append({
                "step_number": i + 1,
                "step_type": step.step_type,
                "status": step.status,
                "description": step.description,
                "completed": step.status == "completed"
            })
    
    # Calculate percentage with more granular tracking
    if current == 0:
        percent = 0
    elif current < total:
        percent = round((current / total) * 100, 1)
    else:
        percent = 100

    snapshot = None
    if req:
        # Helper: build preview and report urls
        def _build_preview_and_reports(r):
            preview = None
            # Prefer existing summary if present
            if getattr(r, 'summary', None):
                preview = str(getattr(r, 'summary'))
            # Otherwise derive from final_result
            if not preview and r.final_result:
                fr = r.final_result
                titles = []
                for a in fr.get('processed_articles', [])[:3]:
                    title = a.get('title') if isinstance(a, dict) else None
                    if title:
                        titles.append(title)
                keywords = fr.get('top_keywords', [])
                if isinstance(keywords, list):
                    kw = ", ".join([k.get('keyword', k) if isinstance(k, dict) else str(k) for k in keywords[:5]])
                else:
                    kw = None
                parts = []
                if titles:
                    parts.append("; ".join(titles))
                if kw:
                    parts.append(f"Keywords: {kw}")
                if parts:
                    preview = " | ".join(parts)
            report_urls = {
                "pdf": f"/research/{r.research_id}/export.pdf",
                "docx": f"/research/{r.research_id}/export.docx",
            }
            return preview, report_urls

        preview, report_urls = _build_preview_and_reports(req)
        # Build workflow_steps
        workflow_steps = []
        for i, step in enumerate(req.steps):
            workflow_steps.append({
                "step_number": i + 1,
                "step_id": step.step_id,
                "step_type": str(step.step_type),
                "description": step.description,
                "status": str(step.status),
                "duration_seconds": step.duration_seconds,
                "timestamp": step.timestamp.isoformat() if step.timestamp else None,
                "output_data": step.output_data,
                "error_message": step.error_message
            })
        # Results
        results = {}
        if req.final_result:
            processed = req.final_result.get("processed_articles", [])
            top_keywords = req.final_result.get("top_keywords", [])
            research_summary = req.final_result.get("research_summary", "")
            results = {
                "processed_articles": processed,
                "top_keywords": top_keywords,
                "research_summary": research_summary,
                "total_articles_processed": len(processed)
            }
            if not results.get("processed_articles") and "analysis" in req.final_result:
                analysis = req.final_result["analysis"]
                results = {
                    "processed_articles": analysis.get("sources", []),
                    "top_keywords": analysis.get("key_findings", []),
                    "total_articles_processed": len(analysis.get("sources", []))
                }
        snapshot = {
            "research_id": req.research_id,
            "topic": req.topic,
            "status": req.status,
            "created_at": req.created_at,
            "completed_at": req.completed_at,
            "trace_log": req.trace_log,
            "workflow_steps": workflow_steps,
            "results": results,
            "steps": [s.model_dump() for s in req.steps],
            "preview": preview,
            "report_urls": report_urls,
        }

    return {
        "research_id": research_id,
        "status": info.get("status", "UNKNOWN"),
        "message": info.get("message", "Processing..."),
        "progress": {
            "current": current, 
            "total": total, 
            "percent": percent,
            "details": progress_details
        },
        "ready": bool(req),
        "snapshot": snapshot,
    }

@app.get("/research/{research_id}", response_model=ResearchResultResponse)
async def get_research_result(research_id: str):
    """
    Get the result of a specific research session
    """
    research_request = research_agent.get_research_request(research_id)
    
    if not research_request:
        raise HTTPException(status_code=404, detail="Research not found")
    
    # Log research request details for monitoring
    logger.info(f"Research request found - ID: {research_request.research_id}, Has results: {research_request.final_result is not None}")
    
    # Convert to response format
    response_data = {
        "research_id": research_request.research_id,
        "topic": research_request.topic,
        "status": research_request.status,
        "created_at": research_request.created_at,
        "completed_at": research_request.completed_at,
        "trace_log": research_request.trace_log,
        "steps": [step.model_dump() for step in research_request.steps],
        "results": research_request.final_result  # Include the research results
    }

    # Derive workflow_steps for UI
    workflow_steps = []
    for i, step in enumerate(research_request.steps):
        workflow_steps.append({
            "step_number": i + 1,
            "step_id": step.step_id,
            "step_type": str(step.step_type),
            "description": step.description,
            "status": str(step.status),
            "duration_seconds": step.duration_seconds,
            "timestamp": step.timestamp.isoformat() if step.timestamp else None,
            "output_data": step.output_data,
            "error_message": step.error_message
        })

    # Build results section from final_result
    results = {}
    if research_request.final_result:
        processed = research_request.final_result.get("processed_articles", [])
        top_keywords = research_request.final_result.get("top_keywords", [])
        results = {
            "processed_articles": processed,
            "top_keywords": top_keywords,
            "total_articles_processed": len(processed)
        }

    # Prefer newer fields; keep legacy analysis fallback
    if not results and research_request.final_result and "analysis" in research_request.final_result:
        analysis = research_request.final_result["analysis"]
        results = {
            "processed_articles": analysis.get("sources", []),
            "top_keywords": analysis.get("key_findings", []),
            "total_articles_processed": len(analysis.get("sources", []))
        }

    # Build preview and report urls
    def _build_preview_and_reports(r):
        preview = None
        if getattr(r, 'summary', None):
            preview = str(getattr(r, 'summary'))
        if not preview and r.final_result:
            fr = r.final_result
            titles = []
            for a in fr.get('processed_articles', [])[:3]:
                title = a.get('title') if isinstance(a, dict) else None
                if title:
                    titles.append(title)
            keywords = fr.get('top_keywords', [])
            if isinstance(keywords, list):
                kw = ", ".join([k.get('keyword', k) if isinstance(k, dict) else str(k) for k in keywords[:5]])
            else:
                kw = None
            parts = []
            if titles:
                parts.append("; ".join(titles))
            if kw:
                parts.append(f"Keywords: {kw}")
            if parts:
                preview = " | ".join(parts)
        report_urls = {
            "pdf": f"/research/{r.research_id}/export.pdf",
            "docx": f"/research/{r.research_id}/export.docx",
        }
        return preview, report_urls

    preview, report_urls = _build_preview_and_reports(research_request)

    # Attach to response
    if workflow_steps:
        response_data["workflow_steps"] = workflow_steps
    if results:
        response_data["results"] = results
    if preview:
        response_data["preview"] = preview
    response_data["report_urls"] = report_urls
    
    return ResearchResultResponse(**response_data)

@app.get("/research", response_model=List[ResearchResultResponse])
async def get_all_research():
    """
    Get all research sessions
    """
    research_requests = research_agent.get_all_research_requests()
    
    results = []
    for request in research_requests:
        response_data = {
            "research_id": request.research_id,
            "topic": request.topic,
            "status": request.status,
            "created_at": request.created_at,
            "completed_at": request.completed_at,
            "trace_log": request.trace_log,
            "steps": [step.model_dump() for step in request.steps]
        }

        # workflow_steps for list view
        workflow_steps = []
        for i, step in enumerate(request.steps):
            workflow_steps.append({
                "step_number": i + 1,
                "step_id": step.step_id,
                "step_type": str(step.step_type),
                "description": step.description,
                "status": str(step.status),
                "duration_seconds": step.duration_seconds,
                "timestamp": step.timestamp.isoformat() if step.timestamp else None,
                "output_data": step.output_data,
                "error_message": step.error_message
            })

        # Build results from final_result
        results_block = {}
        if request.final_result:
            processed = request.final_result.get("processed_articles", [])
            top_keywords = request.final_result.get("top_keywords", [])
            results_block = {
                "processed_articles": processed,
                "top_keywords": top_keywords,
                "total_articles_processed": len(processed),
                "research_summary": request.final_result.get("research_summary", "")
            }
        if not results_block and request.final_result and "analysis" in request.final_result:
            analysis = request.final_result["analysis"]
            results_block = {
                "processed_articles": analysis.get("sources", []),
                "top_keywords": analysis.get("key_findings", []),
                "total_articles_processed": len(analysis.get("sources", [])),
                "research_summary": analysis.get("research_summary", "")
            }

        if workflow_steps:
            response_data["workflow_steps"] = workflow_steps
        if results_block:
            response_data["results"] = results_block
        
        results.append(ResearchResultResponse(**response_data))
    
    return results

@app.delete("/research/{research_id}")
async def delete_research(research_id: str):
    """
    Delete a specific research session
    """
    from database import db_manager
    deleted = db_manager.delete_research_request(research_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="Research not found")
    # Drop transient job status cache
    if research_id in job_status:
        del job_status[research_id]
    return {"deleted": True, "research_id": research_id}

@app.delete("/research")
async def delete_all_research():
    """
    Delete all research sessions
    """
    from database import db_manager
    count = db_manager.delete_all_research_requests()
    job_status.clear()
    return {"deleted": count}

@app.options("/{path:path}")
async def options_handler(path: str):
    """
    Handle OPTIONS requests for CORS
    """
    return {"message": "OK"}

@app.get("/health")
async def health_check():
    """
    Health check endpoint
    """
    return {"status": "healthy", "timestamp": datetime.now(timezone.utc)}

# Utility to build a simple PDF
def _build_pdf_for_research(req: ResearchRequest) -> BytesIO:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        ListFlowable,
        ListItem,
        Table,
        TableStyle,
        PageBreak,
    )
    from reportlab.lib.units import inch
    from reportlab.lib import colors

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        title=f"Research - {req.topic}",
        leftMargin=0.75 * inch,
        rightMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Small", parent=styles["Normal"], fontSize=9, leading=11))
    styles.add(ParagraphStyle(name="Muted", parent=styles["Normal"], textColor=colors.gray))
    styles.add(ParagraphStyle(name="H3", parent=styles["Heading2"], fontSize=12, leading=14))

    elems = []
    # Simple HTML/tag sanitizer for ReportLab paragraphs
    import re as _re
    import html as _html

    def _clean_text(value) -> str:
        s = str(value) if value is not None else ""
        s = _html.unescape(s)
        s = _re.sub(r"</?span[^>]*>", "", s, flags=_re.IGNORECASE)
        s = _re.sub(r"<[^>]*>", "", s)
        s = s.replace("<", "").replace(">", "")
        s = _re.sub(r"\s+", " ", s).strip()
        return s

    def _coerce_keywords(kws):
        out = []
        for k in kws or []:
            if isinstance(k, dict):
                out.append(_clean_text(k.get("keyword") or k.get("text") or ""))
            else:
                out.append(_clean_text(k))
        return [k for k in out if k]

    # Header
    elems.append(Paragraph(_clean_text(req.topic), styles['Title']))
    meta = [
        f"Status: {_clean_text(req.status)}",
        f"Created: {_clean_text(req.created_at)}",
    ]
    if req.completed_at:
        meta.append(f"Completed: {_clean_text(req.completed_at)}")
    elems.append(Paragraph(" | ".join(meta), styles['Muted']))
    elems.append(Spacer(1, 0.25 * inch))

    # Summary
    elems.append(Paragraph("Summary", styles['Heading2']))
    summary_text = None
    if getattr(req, 'summary', None):
        summary_text = _clean_text(getattr(req, 'summary'))
    elif req.final_result and req.final_result.get("research_summary"):
        # Use AI-generated research summary
        summary_text = _clean_text(req.final_result.get("research_summary"))
    elif req.final_result:
        # Derive a compact summary if not explicitly set
        titles = []
        for a in (req.final_result.get("processed_articles", []) or [])[:3]:
            title = a.get("title") if isinstance(a, dict) else None
            if title:
                titles.append(_clean_text(title))
        kw = ", ".join(_coerce_keywords(req.final_result.get("top_keywords")))
        parts = []
        if titles:
            parts.append("; ".join(titles))
        if kw:
            parts.append(f"Top keywords: {kw}")
        if parts:
            summary_text = " | ".join(parts)
    elems.append(Paragraph(_clean_text(summary_text or "No summary available."), styles['Normal']))
    elems.append(Spacer(1, 0.2 * inch))

    # Keywords
    kws = _coerce_keywords(req.final_result.get("top_keywords") if getattr(req, 'final_result', None) else [])
    if kws:
        elems.append(Paragraph("Top Keywords", styles['H3']))
        kw_items = [ListItem(Paragraph(_clean_text(k), styles['Normal'])) for k in kws]
        elems.append(ListFlowable(kw_items, bulletType='bullet'))
        elems.append(Spacer(1, 0.2 * inch))

    # Workflow steps
    elems.append(Paragraph("Workflow Steps", styles['Heading2']))
    if req.steps:
        step_items = []
        for i, s in enumerate(req.steps, 1):
            step_text = f"{i}. {s.step_type} - {s.status}: {s.description}"
            step_items.append(ListItem(Paragraph(_clean_text(step_text), styles['Normal'])))
        elems.append(ListFlowable(step_items, bulletType='1'))
    else:
        elems.append(Paragraph("No steps recorded.", styles['Small']))
    elems.append(Spacer(1, 0.25 * inch))

    # Articles table
    elems.append(Paragraph("Sources Reviewed", styles['Heading2']))
    processed = (req.final_result or {}).get("processed_articles", []) if getattr(req, 'final_result', None) else []
    if processed:
        data = [["Title", "Source", "URL", "Summary"]]
        for art in processed:
            if isinstance(art, dict):
                title = _clean_text(art.get("title") or "-")
                source = _clean_text(art.get("source") or art.get("site") or "-")
                url = _clean_text(art.get("url") or art.get("link") or "-")
                snippet = _clean_text(art.get("summary") or art.get("snippet") or art.get("description") or "-")
            else:
                title = _clean_text(str(art))
                source = "-"
                url = "-"
                snippet = "-"
            # Limit extremely long fields for PDF layout
            def _truncate(t, n=180):
                return (t[: n - 1] + "…") if len(t) > n else t
            data.append([_truncate(title, 80), _truncate(source, 30), _truncate(url, 80), _truncate(snippet, 180)])

        table = Table(data, repeatRows=1, colWidths=[2.3 * inch, 1.2 * inch, 2.3 * inch, 2.7 * inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#cccccc')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#fbfbfb')]),
        ]))
        elems.append(table)
    else:
        elems.append(Paragraph("No sources found.", styles['Small']))

    elems.append(Spacer(1, 0.25 * inch))

    # Footer with page numbers
    def _add_page_number(canvas, _doc):
        page_num = canvas.getPageNumber()
        canvas.setFont("Helvetica", 9)
        canvas.setFillColor(colors.gray)
        canvas.drawRightString(letter[0] - 0.75 * inch, 0.5 * inch, f"Page {page_num}")

    doc.build(elems, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
    buffer.seek(0)
    return buffer

# Utility to build a DOCX
def _build_docx_for_research(req: ResearchRequest) -> BytesIO:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.enum.text import WD_BREAK

    # Sanitizer
    import re as _re
    import html as _html
    def _clean_text(value) -> str:
        s = str(value) if value is not None else ""
        s = _html.unescape(s)
        s = _re.sub(r"</?span[^>]*>", "", s, flags=_re.IGNORECASE)
        s = _re.sub(r"<[^>]*>", "", s)
        s = s.replace("<", "").replace(">", "")
        s = _re.sub(r"\s+", " ", s).strip()
        return s

    def _coerce_keywords(kws):
        out = []
        for k in kws or []:
            if isinstance(k, dict):
                out.append(_clean_text(k.get("keyword") or k.get("text") or ""))
            else:
                out.append(_clean_text(k))
        return [k for k in out if k]

    doc = Document()
    doc.core_properties.title = f"Research - {req.topic}"
    doc.add_heading(f"{_clean_text(req.topic)}", 0)

    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Status: {_clean_text(req.status)}\n").font.size = Pt(10)
    meta_p.add_run(f"Created: {_clean_text(req.created_at)}\n").font.size = Pt(10)
    if req.completed_at:
        meta_p.add_run(f"Completed: {_clean_text(req.completed_at)}\n").font.size = Pt(10)

    # Summary
    doc.add_heading("Summary", level=1)
    summary_text = None
    if getattr(req, 'summary', None):
        summary_text = _clean_text(getattr(req, 'summary'))
    elif getattr(req, 'final_result', None) and req.final_result.get("research_summary"):
        # Use AI-generated research summary
        summary_text = _clean_text(req.final_result.get("research_summary"))
    elif getattr(req, 'final_result', None):
        titles = []
        for a in (req.final_result.get("processed_articles", []) or [])[:3]:
            title = a.get("title") if isinstance(a, dict) else None
            if title:
                titles.append(_clean_text(title))
        kw = ", ".join(_coerce_keywords(req.final_result.get("top_keywords")))
        parts = []
        if titles:
            parts.append("; ".join(titles))
        if kw:
            parts.append(f"Top keywords: {kw}")
        if parts:
            summary_text = " | ".join(parts)
    doc.add_paragraph(_clean_text(summary_text or "No summary available."))

    # Keywords
    kws = _coerce_keywords(req.final_result.get("top_keywords") if getattr(req, 'final_result', None) else [])
    if kws:
        doc.add_heading("Top Keywords", level=2)
        for k in kws:
            doc.add_paragraph(_clean_text(k), style='List Bullet')

    # Workflow steps
    doc.add_heading("Workflow Steps", level=1)
    if req.steps:
        for i, s in enumerate(req.steps, 1):
            doc.add_paragraph(f"{i}. {s.step_type} - {s.status}: {s.description}", style='List Number')
    else:
        doc.add_paragraph("No steps recorded.")

    # Sources table
    doc.add_heading("Sources Reviewed", level=1)
    processed = (req.final_result or {}).get("processed_articles", []) if getattr(req, 'final_result', None) else []
    if processed:
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Title"
        hdr_cells[1].text = "Source"
        hdr_cells[2].text = "URL"
        hdr_cells[3].text = "Summary"
        for art in processed:
            if isinstance(art, dict):
                title = _clean_text(art.get("title") or "-")
                source = _clean_text(art.get("source") or art.get("site") or "-")
                url = _clean_text(art.get("url") or art.get("link") or "-")
                snippet = _clean_text(art.get("summary") or art.get("snippet") or art.get("description") or "-")
            else:
                title, source, url, snippet = _clean_text(str(art)), "-", "-", "-"
            row_cells = table.add_row().cells
            row_cells[0].text = title
            row_cells[1].text = source
            row_cells[2].text = url
            row_cells[3].text = snippet
    else:
        doc.add_paragraph("No sources found.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.get("/research/{research_id}/export.pdf")
async def export_research_pdf(research_id: str):
    req = research_agent.get_research_request(research_id)
    if not req:
        raise HTTPException(status_code=404, detail="Research not found")
    try:
        pdf = _build_pdf_for_research(req)
    except ModuleNotFoundError as e:
        # ReportLab not installed
        raise HTTPException(
            status_code=501,
            detail="PDF export requires the 'reportlab' package. Install with: pip install reportlab"
        ) from e
    return StreamingResponse(pdf, media_type="application/pdf", headers={
        "Content-Disposition": f"attachment; filename=research_{research_id}.pdf"
    })

@app.get("/research/{research_id}/export.docx")
async def export_research_docx(research_id: str):
    req = research_agent.get_research_request(research_id)
    if not req:
        raise HTTPException(status_code=404, detail="Research not found")
    try:
        docx_io = _build_docx_for_research(req)
    except ModuleNotFoundError as e:
        # python-docx not installed
        raise HTTPException(
            status_code=501,
            detail="DOCX export requires the 'python-docx' package. Install with: pip install python-docx"
        ) from e
    return StreamingResponse(docx_io, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={
        "Content-Disposition": f"attachment; filename=research_{research_id}.docx"
    })

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)

